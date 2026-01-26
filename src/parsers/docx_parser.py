from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from .base_parser import BaseParser


class DOCXParser(BaseParser):
    """Parser for DOCX documents"""

    def parse(self) -> str:
        """Extract content from DOCX and convert to markdown"""
        markdown_lines = []
        markdown_lines.append(f"# {self.file_path.stem}\n")
        markdown_lines.append(f"*Source: {self.file_path.name}*\n")
        markdown_lines.append("---\n")

        try:
            doc = Document(self.file_path)

            # Process each element in the document
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    # Paragraph
                    para = self._find_paragraph(doc, element)
                    if para:
                        markdown_lines.append(self._paragraph_to_markdown(para))
                elif element.tag.endswith('tbl'):
                    # Table
                    table = self._find_table(doc, element)
                    if table:
                        markdown_lines.append(self._table_to_markdown(table))
                        markdown_lines.append("\n")

        except Exception as e:
            markdown_lines.append(f"\n**Error parsing DOCX:** {str(e)}\n")

        return "\n".join(markdown_lines)

    def _find_paragraph(self, doc, element):
        """Find paragraph object from element"""
        for para in doc.paragraphs:
            if para._element == element:
                return para
        return None

    def _find_table(self, doc, element):
        """Find table object from element"""
        for table in doc.tables:
            if table._element == element:
                return table
        return None

    def _paragraph_to_markdown(self, para: Paragraph) -> str:
        """Convert paragraph to markdown format"""
        text = para.text.strip()
        if not text:
            return ""

        # Detect heading levels based on style
        style_name = para.style.name.lower()
        if 'heading 1' in style_name:
            return f"# {text}\n"
        elif 'heading 2' in style_name:
            return f"## {text}\n"
        elif 'heading 3' in style_name:
            return f"### {text}\n"
        elif 'heading 4' in style_name:
            return f"#### {text}\n"
        elif 'heading 5' in style_name:
            return f"##### {text}\n"
        elif 'heading 6' in style_name:
            return f"###### {text}\n"
        else:
            # Check for bold text
            if para.runs:
                all_bold = all(run.bold for run in para.runs if run.text.strip())
                if all_bold and text:
                    return f"**{text}**\n"
            return f"{text}\n"

    def _table_to_markdown(self, table: Table) -> str:
        """Convert table to markdown format"""
        if not table.rows:
            return ""

        markdown_table = []

        # Header row
        header_cells = table.rows[0].cells
        header = [cell.text.strip() for cell in header_cells]
        markdown_table.append("| " + " | ".join(header) + " |")
        markdown_table.append("| " + " | ".join(["---"] * len(header)) + " |")

        # Data rows
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            markdown_table.append("| " + " | ".join(cells) + " |")

        return "\n".join(markdown_table)
